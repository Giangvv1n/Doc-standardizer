from __future__ import annotations

import os
from pathlib import Path
import pytest
import docx

from app.ingestion.reader import DocxReader
from app.segmentation.segmenter import DocumentSegmenter, DocumentSection
from app.matching.matcher import SectionMatcher
from app.template.renderer import DocxRenderer


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> Path:
    """Creates a temporary sample docx file for unit testing."""
    file_path = tmp_path / "test_doc.docx"
    doc = docx.Document()
    doc.add_heading("Tiêu đề kiểm thử", level=0)
    
    doc.add_heading("Giới thiệu chung", level=1)
    doc.add_paragraph("Nội dung giới thiệu doanh nghiệp.")
    
    doc.add_heading("Mục đích dự án", level=1)
    doc.add_paragraph("Nội dung mục tiêu dự án.")
    
    doc.save(str(file_path))
    return file_path


def test_reader(sample_docx_path: Path) -> None:
    """Verifies that DocxReader correctly parses elements in sequential order."""
    reader = DocxReader()
    blocks = reader.read(str(sample_docx_path))
    
    # Check that blocks contains paragraph/headings
    assert len(blocks) > 0
    # The first text item should be the main title
    assert hasattr(blocks[0], "text")
    assert "Tiêu đề kiểm thử" in blocks[0].text


def test_segmenter(sample_docx_path: Path) -> None:
    """Verifies that DocumentSegmenter splits document blocks into sections correctly."""
    reader = DocxReader()
    blocks = reader.read(str(sample_docx_path))
    
    segmenter = DocumentSegmenter()
    sections = segmenter.segment(blocks)
    
    # We should have at least 2 structured sections
    assert len(sections) >= 2
    headings = [sec.heading_text for sec in sections]
    assert "Giới thiệu chung" in headings
    assert "Mục đích dự án" in headings


def test_matcher() -> None:
    """Verifies that SectionMatcher maps synonyms and headings to canonical sections correctly."""
    matcher = SectionMatcher()
    
    # Test section matching for a clear synonym
    sec = DocumentSection(heading_text="Về chúng tôi", heading_level=1)
    canonical, score, status = matcher.match_section(sec)
    
    assert canonical == "Thông tin doanh nghiệp"
    assert score >= 0.70
    assert status in ("auto_accepted", "needs_review")


def test_renderer_and_end_to_end(sample_docx_path: Path, tmp_path: Path) -> None:
    """Verifies the rendering phase and verifies that a standard output is produced."""
    reader = DocxReader()
    blocks = reader.read(str(sample_docx_path))
    
    segmenter = DocumentSegmenter()
    sections = segmenter.segment(blocks)
    
    matcher = SectionMatcher()
    mappings = matcher.match_document(sections)
    
    # Render with custom output settings overrides
    renderer = DocxRenderer()
    
    # Temporarily override output directory for testing
    from app.config import settings
    original_output_dir = settings.output_dir
    settings.output_dir = tmp_path
    
    try:
        output_path = renderer.render(
            mappings=mappings,
            output_file_name="test_std_output.docx"
        )
        
        assert output_path.exists()
        assert output_path.stat().st_size > 0
        
        # Verify that output file contains the standardized text
        out_doc = docx.Document(str(output_path))
        text_content = "\n".join(p.text for p in out_doc.paragraphs)
        assert "Nội dung giới thiệu doanh nghiệp" in text_content
        assert "Nội dung mục tiêu dự án" in text_content
    finally:
        settings.output_dir = original_output_dir
