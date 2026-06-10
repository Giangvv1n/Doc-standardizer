from __future__ import annotations

from copy import deepcopy
import json
from pathlib import Path
from docxtpl import DocxTemplate
from docx.text.paragraph import Paragraph
from docx.table import Table
from loguru import logger

from app.config import settings
from app.constants import DEFAULT_SECTION_PLACEHOLDERS
from app.segmentation.segmenter import DocumentSection


import re
import docx


class DocxRenderer:
    """Renders standardized word documents by inserting matching section blocks into a standard template."""

    def __init__(self, template_path: Path | None = None) -> None:
        self.template_path = template_path if template_path is not None else settings.template_path
        self.placeholder_mapping = self._load_placeholder_mapping()

    @staticmethod
    def extract_template_schema(template_path: Path) -> dict[str, str]:
        """
        Scans the template document to find placeholders of type {{p placeholder}} or {{placeholder}}
        and maps them to their closest preceding heading text.
        """
        if not template_path.exists():
            logger.warning(f"Template path {template_path} does not exist. Cannot extract dynamic schema.")
            return {}

        try:
            doc = docx.Document(str(template_path))
            placeholder_pattern = re.compile(r"\{\{\s*(?:p\s+)?(\w+)\s*\}\}")
            
            mapping = {}
            current_heading = None
            
            from app.constants import COMMON_HEADING_STYLES, HEADING_NUMBER_PATTERNS
            heading_regexes = [re.compile(p) for p in HEADING_NUMBER_PATTERNS]
            
            for p in doc.paragraphs:
                text = p.text.strip()
                if not text:
                    continue
                
                # Detect if paragraph is a heading
                style_name = p.style.name if p.style else ""
                is_heading = False
                
                if style_name in COMMON_HEADING_STYLES:
                    is_heading = True
                elif len(text) <= 120 and len(text.split()) <= 15:
                    for regex in heading_regexes:
                        if regex.match(text):
                            is_heading = True
                            break
                    
                    runs = [r for r in p.runs if r.text.strip()]
                    if runs and all(r.bold for r in runs):
                        is_heading = True
                        
                if is_heading:
                    current_heading = text
                    
                # Check for placeholder
                match = placeholder_pattern.search(text)
                if match:
                    placeholder_var = match.group(1)
                    sec_name = current_heading if current_heading else placeholder_var.replace("_", " ").title()
                    mapping[sec_name] = placeholder_var
                    current_heading = None
                    
            logger.info(f"Dynamically extracted {len(mapping)} placeholder mappings from template.")
            return mapping
        except Exception as e:
            logger.error(f"Failed to dynamically extract template schema from {template_path}: {e}")
            return {}

    def _load_placeholder_mapping(self) -> dict[str, str]:
        """Loads section-to-placeholder mappings dynamically from the template, falling back to schema JSON or default constants."""
        # Try dynamic extraction first
        dynamic_mapping = self.extract_template_schema(self.template_path)
        if dynamic_mapping:
            return dynamic_mapping

        # Fallback 1: section_schema.json
        schema_path = settings.section_schema_path
        if schema_path.exists():
            try:
                logger.info(f"Loading section mapping from {schema_path}")
                with open(schema_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    mapping = {}
                    for section in data.get("sections", []):
                        mapping[section["name"]] = section["placeholder"]
                    return mapping
            except Exception as e:
                logger.error(f"Error reading placeholder mapping from schema {schema_path}: {e}")

        # Fallback 2: Default constants
        logger.info("Using default placeholder mapping.")
        return dict(DEFAULT_SECTION_PLACEHOLDERS)

    def _copy_block_to_subdoc(self, subdoc: Any, block: Paragraph | Table) -> None:
        """Copies a paragraph or table block from a source document into a docxtpl SubDoc."""
        if isinstance(block, Paragraph):
            # Add a dummy paragraph to get the parent container in SubDoc
            dummy_p = subdoc.add_paragraph()
            parent = dummy_p._p.getparent()
            
            # Deepcopy the source paragraph XML and replace the dummy paragraph element
            copied_p_xml = deepcopy(block._p)
            parent.replace(dummy_p._p, copied_p_xml)
        elif isinstance(block, Table):
            # Table XML can be directly appended to the SubDoc body
            copied_tbl_xml = deepcopy(block._tbl)
            subdoc.element.body.append(copied_tbl_xml)

    def render(
        self,
        mappings: dict[str, list[tuple[DocumentSection, float, str]]],
        llm_extractions: dict[str, str] | None = None,
        output_file_name: str = "standardized_output.docx",
    ) -> Path:
        """
        Renders matched content and LLM-extracted content into the template.
        Saves output in the output directory.
        """
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template document not found at {self.template_path}")

        logger.info(f"Loading template: {self.template_path}")
        doc = DocxTemplate(str(self.template_path))
        context: dict[str, Any] = {}

        # Process each canonical section mapping
        for canonical_name, placeholder in self.placeholder_mapping.items():
            matched_items = mappings.get(canonical_name, [])

            if matched_items:
                logger.info(f"Rendering rich content for section: '{canonical_name}' -> placeholder: '{placeholder}'")
                subdoc = doc.new_subdoc()
                
                # Append each matching section block-by-block
                for section, _, _ in matched_items:
                    for block in section.blocks:
                        self._copy_block_to_subdoc(subdoc, block)
                
                context[placeholder] = subdoc

            elif llm_extractions and canonical_name in llm_extractions:
                logger.info(f"Rendering LLM-extracted text for section: '{canonical_name}' -> placeholder: '{placeholder}'")
                subdoc = doc.new_subdoc()
                text = llm_extractions[canonical_name]
                subdoc.add_paragraph(text)
                context[placeholder] = subdoc

            else:
                logger.info(f"No content found for section: '{canonical_name}', leaving placeholder blank.")
                # We can either leave it blank or insert a placeholder notice
                context[placeholder] = ""

        # Ensure output directory exists
        settings.output_dir.mkdir(parents=True, exist_ok=True)
        output_path = settings.output_dir / output_file_name

        logger.info(f"Saving rendered document to {output_path}")
        doc.render(context)
        doc.save(str(output_path))
        return output_path
